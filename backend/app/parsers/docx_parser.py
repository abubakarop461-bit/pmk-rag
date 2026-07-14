import docx
from typing import List
from langchain_core.documents import Document
from app.parsers.base_parser import BaseParser
from loguru import logger

class DocxParser(BaseParser):
    def parse(self, file_path: str) -> List[Document]:
        """
        Loads DOCX file using python-docx and extracts paragraph text layers.
        """
        logger.info(f"DocxParser starting extraction on: {file_path}")
        try:
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            text = "\n".join(full_text).strip()
            logger.info(f"DocxParser successfully extracted content from {file_path}.")
            return [
                Document(
                    page_content=text,
                    metadata={
                        "source": file_path,
                        "page": 1
                    }
                )
            ]
        except Exception as e:
            logger.error(f"python-docx failed to parse Word file {file_path}: {e}")
            raise RuntimeError(f"Word parsing failed: {e}")
